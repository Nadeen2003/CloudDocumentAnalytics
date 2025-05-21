import fitz  # PyMuPDF
import docx
import re
import os
import time
from typing import Tuple, List, Dict

def extract_title_from_pdf(file_path):
    """Extract title from PDF document."""
    doc = fitz.open(file_path)
    for page in doc:
        text = page.get_text().strip()
        lines = text.split('\n')
        for line in lines:
            if len(line.strip()) > 10:
                return line.strip()
    return "No Title"

def extract_title_from_docx(file_path):
    """Extract title from DOCX document."""
    doc = docx.Document(file_path)
    for para in doc.paragraphs:
        if para.text.strip():
            return para.text.strip()
    return "No Title"

def search_text_in_file(file_path: str, keyword: str) -> Tuple[bool, str, List[Dict], str]:
    """
    Search for text in a document and return matches information and the keyword.
    
    Args:
        file_path (str): Path to the document
        keyword (str): Text to search for
    
    Returns:
        Tuple[bool, str, List[Dict], str]: (found, full_text, matches, keyword)
            found: Whether the keyword was found
            full_text: The full text of the document
            matches: List of dictionaries containing match information
            keyword: The original search keyword
    """
    start_time = time.time()
    matches = []
    
    print(f"Searching for: {keyword}")  # Debug print
    
    try:
        # We will still use regex to find initial matches, but highlighting will be redone based on the text
        # Use a pattern that looks for the keyword surrounded by non-word characters or start/end of string
        # Escape the keyword first to handle special regex characters
        escaped_keyword = re.escape(keyword)
        search_pattern = r'(?:^|\W)' + escaped_keyword + r'(?:\W|$)'
        print(f"Using pattern: {search_pattern}")  # Debug print

        if file_path.endswith('.pdf'):
            doc = fitz.open(file_path)
            full_text = ""
            for page_num, page in enumerate(doc):
                text = page.get_text()
                full_text += text
                # Find matches with their positions
                for match in re.finditer(search_pattern, text, re.IGNORECASE):
                    # Store basic match info. Exact highlighting will be done in highlight_text
                    matches.append({
                        'page': page_num + 1,
                        'start': match.start(), # Keep original positions for context if needed
                        'end': match.end(),
                        'text': match.group()
                    })
        elif file_path.endswith('.docx'):
            doc = docx.Document(file_path)
            full_text = ""
            for para_num, para in enumerate(doc.paragraphs):
                text = para.text + "\n"
                full_text += text
                # Find matches with their positions
                for match in re.finditer(search_pattern, text, re.IGNORECASE):
                    # Store basic match info. Exact highlighting will be done in highlight_text
                    matches.append({
                        'paragraph': para_num + 1,
                        'start': match.start(), # Keep original positions for context if needed
                        'end': match.end(),
                        'text': match.group()
                    })
        
        search_time = time.time() - start_time
        print(f"Search completed in {search_time:.2f} seconds, {len(matches)} matches found.")  # Debug print
        
        return len(matches) > 0, full_text, matches, keyword
    
    except Exception as e:
        print(f"Error searching file: {str(e)}")
        return False, "", [], keyword

def highlight_text(text: str, keyword: str) -> str:
    """
    Highlight occurrences of the keyword in the text using HTML span with background color.
    Finds occurrences directly in the provided text.
    
    Args:
        text (str): The full text of the document or a relevant snippet.
        keyword (str): The keyword to highlight.
    
    Returns:
        str: Text with highlighted keyword occurrences using HTML
    """
    if not keyword:
        return text
    
    # Use regex to find occurrences of the keyword directly in the text to be highlighted
    # Use word boundaries to avoid partial matches within words
    # Escape keyword for regex safety
    escaped_keyword = re.escape(keyword)
    search_pattern = r'\b' + escaped_keyword + r'\b'

    highlighted_text = ""
    last_end = 0
    
    for match in re.finditer(search_pattern, text, re.IGNORECASE):
        start, end = match.span()
        # Append text before the match, the highlighted match, and update last_end
        highlighted_text += text[last_end:start]
        highlighted_text += f"<span style=\"background-color: #ADD8E6;\">{text[start:end]}</span>"
        last_end = end
    
    # Append any remaining text after the last match
    highlighted_text += text[last_end:]
    
    return highlighted_text
